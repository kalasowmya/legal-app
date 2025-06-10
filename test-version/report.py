# Import necessary libraries
from docx import Document  # For creating and editing Word documents
from docx.shared import Inches  # (Not used here but typically for setting dimensions)
from docx.oxml.ns import qn  # (Not used here but typically for XML namespaces)
from bs4 import BeautifulSoup  # For parsing HTML
from io import BytesIO  # To create an in-memory binary stream

# Main function that converts risk analysis HTML into a downloadable .docx file
def report_downloader(risk_analysis_html: str,logo_path: str) -> bytes:
    # Parse the HTML input using BeautifulSoup
    soup = BeautifulSoup(risk_analysis_html, "html.parser")
    
    # Create a new Word document
    doc = Document()

    # Insert the FM logo at the top-left corner
    doc.sections[0].header.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1))
    
    # Helper function to add formatted text (runs) to a paragraph
    def add_formatted_run(paragraph, node):
        """Adds text with formatting from HTML tags."""
        
        # If the node is plain text, add it directly
        if isinstance(node, str):
            paragraph.add_run(node)
        
        # If it's a bold tag, add bold text
        elif node.name == 'strong' or node.name == 'b':
            run = paragraph.add_run(node.get_text())
            run.bold = True

        # If it's a strikethrough tag, apply strike formatting
        elif node.name == 'del':
            run = paragraph.add_run(node.get_text())
            run.font.strike = True

        # If it's italicized text, apply italic formatting
        elif node.name == 'em' or node.name == 'i':
            run = paragraph.add_run(node.get_text())
            run.italic = True

        else:
            # If it's another tag, process its children recursively
            for sub in node.contents:
                add_formatted_run(paragraph, sub)

    # Iterate through each top-level element in the HTML content
    for element in soup.contents:
        
        # If the element is a paragraph
        if element.name == "p":
            para = doc.add_paragraph()  # Add a new paragraph to the doc
            for node in element.contents:  # Go through its child elements
                add_formatted_run(para, node)  # Add formatted text accordingly

        # If the element is a table
        elif element.name == "table":
            rows = element.find_all("tr")  # Get all rows in the table
            if not rows:
                continue  # Skip if the table is empty

            # Create a new Word table with same number of rows and columns
            table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))
            table.style = 'Table Grid'  # Apply a grid style to the table

            # Loop over each row and cell to populate the Word table
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])  # Get all cells in the row
                for j, cell in enumerate(cells):
                    para = table.cell(i, j).paragraphs[0]  # Get the paragraph in the cell
                    for node in cell.contents:  # Process each part of the cell
                        add_formatted_run(para, node)

    # Create an in-memory buffer to hold the .docx file
    buffer = BytesIO()
    doc.save(buffer)  # Save the Word document into the buffer
    buffer.seek(0)  # Move to the start of the buffer

    return buffer.getvalue()  # Return the bytes representation of the Word file
