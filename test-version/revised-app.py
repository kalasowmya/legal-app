import streamlit as st
import fitz  # PyMuPDF
import docx
import re
import markdown
# from langchain.chat_models import ChatGoogleGenerativeAI
from langchain.schema import HumanMessage, SystemMessage
# from report_downloader import report_downloader  # assuming your custom function
from langchain_core.prompts import ChatPromptTemplate
# from langchain_google_genai import google_configure
import streamlit as st
import fitz  # PyMuPDF for extracting text from PDFs
import docx  # For Word document support
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import SystemMessage, HumanMessage
import legal_prompt
# import instructions
# import instructions_updated as instructions
from langchain.text_splitter import RecursiveCharacterTextSplitter
from google.generativeai import configure as google_configure
# import keys
from io import BytesIO
import markdown
import re
import report
import io
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes

# Configure Gemini
google_configure(api_key=st.secrets.GOOGLE_API_KEY)
gemini_model = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro", 
    temperature=0.2,# deterministic behavior
    max_output_tokens= 8192 
)

# gemini_model = ChatGoogleGenerativeAI(
#     # model="gemini-1.5-flash",
#     model="gemini-1.5-pro" ,
#     google_api_key=keys.GOOGLE_API_KEY,
#     temperature=0.2,# deterministic behavior
#     max_output_tokens= 8192 
# )

# Constants
GEMINI_MAX_WORDS = 150000  # approximate large limit, adjust if needed

# # Extract PDF text
# def extract_text_from_pdf(pdf_file):
#     doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
#     return "\n".join([page.get_text("text") for page in doc])

# # Extract DOCX text
# def extract_text_from_word(word_file):
#     doc = docx.Document(word_file)
#     return "\n".join([para.text for para in doc.paragraphs])

#adding ocr for scanned docs

# Extract text from PDF (including scanned with OCR)
def extract_text_from_pdf(pdf_file):
    text = ""
    pdf_bytes = pdf_file.read()
    
    # Try extracting selectable text using PyMuPDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        page_text = page.get_text("text")
        text += page_text + "\n"
    
    # If the text is mostly empty, apply OCR
    if len(text.strip()) < 50:
        st.info("Applying OCR for scanned contracts...")
        # print("Applying OCR for scanned PDF...")
        images = convert_from_bytes(pdf_bytes)
        text = "\n".join([pytesseract.image_to_string(img) for img in images])
    
    return text

# Extract text from DOCX
def extract_text_from_word(word_file):
    doc = docx.Document(word_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # If no text found (possibly a scanned DOCX with images), apply OCR
    if len(text.strip()) < 50:
        st.info("Applying OCR for scanned contracts...")
        # print("Applying OCR for scanned DOCX...")
        for shape in doc.inline_shapes:
            try:
                image_bytes = shape.image.blob
                image = Image.open(io.BytesIO(image_bytes))
                text += pytesseract.image_to_string(image) + "\n"
            except Exception as e:
                print(f"OCR failed on shape: {e}")
    
    return text

# Chunk text based on word count
def split_text(text, chunk_size):
    words = text.split()
    return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]

# Analyze contract using Gemini
def analyze_contract(text, rulebook, instructions):
    messages = [
        SystemMessage(content=instructions),
        SystemMessage(content=rulebook),
        HumanMessage(content=f"Analyze the following contract text:\n{text}")
    ]
    response = gemini_model.invoke(messages)
    return response.content if response else "No response from Gemini."

# Default instructions (edit as needed)
DEFAULT_INSTRUCTIONS = """
You are a legal assistant specialized in Indian domestic contracts.
You have to create a **Deviation Sheet** in the following tabular format as a response.

### Deviation Sheet Format (Mandatory):
Column Names: 
●Original Clause Number and Reference 
●Original Clause
●Revised Clause
●Risk Summary

| Original Clause Number and Reference | Original Clause             | Revised Clause                         | Risk Summary                                |
|----------------------------------    |-----------------------------|-------------------------------------   |-------------------------------------------- |
| Clause Number or Reference           | Original Clause & Reference | Modified clause(strikethrough and bold)| Explanation of risks and reasons for changes|


**Rules for Deviation Sheet:**
- DO NOT skip this table.
- DO NOT add missing clauses to the Deviation Sheet. Log them separately along with the risk summary, in "Other Clauses to be looked at:" below the sheet. 
  The Deviation Sheet is only for deviations from existing clauses.
- Only include in the Deviation Sheet those clauses that are present in the contract and have identified legal or commercial risks.
- Only include the specific sub-clause (e.g., 17.2.3) where the legal or commercial risk exists.
- DO NOT use placeholders such as “The entire Clause”or "Entire Clause" or “Refer Clause” in the deviation sheet,instead mention actual clause.
- In “Revised Clause”, STRICTLY apply:
    - ~~strikethrough~~ for deletions.
    - **bold** for additions.
- DO NOT paraphrase or improve clauses unless there is a clear legal risk.
- Ensure that no changes are made to statements that convey the same meaning, even if they differ grammatically.
- Ensure that all references, definitions, and clauses align with the contract’s existing terms and structure. 
- Display all the modified clauses in the Deviation Sheet ONLY.

"""

# Rulebook default (replace with your real rulebook content)
DEFAULT_RULEBOOK = legal_prompt.rulebook

# Streamlit UI
st.set_page_config(page_title="LEGAL CONTRACT REVIEW: DOMESTIC ORDERS (INDIA)", layout="wide")

# Sidebar
with st.sidebar:
    st.image("fm-logo.png", use_container_width=True)
    st.write("Upload your Domestic Order Contracts (India) and receive an instant Legal Risk Assessment.")
    uploaded_file = st.file_uploader("Upload your contract (PDF or Word)", type=["pdf", "docx"], key="contract")

st.header("Legal Risk Assessment Tool — Version 1.0 - Revised (test)")

# Editable Instructions
st.subheader("Modify Default Instructions")
instructions_text = st.text_area(
    "Modify the default instructions as needed:", 
    st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS), 
    height=300
)
if st.button("Save Instructions"):
    st.session_state["saved_instructions"] = instructions_text
    st.success("Instructions saved.")

# Editable Rulebook
st.subheader("Modify Rulebook")
rulebook_text = st.text_area(
    "Modify the rulebook as needed:", 
    st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK), 
    height=500
)
if st.button("Save Rulebook"):
    st.session_state["saved_rulebook"] = rulebook_text
    st.success("Rulebook saved.")

# Load saved instruction/rulebook
instructions_text = st.session_state.get("saved_instructions", DEFAULT_INSTRUCTIONS)
rulebook_text = st.session_state.get("saved_rulebook", DEFAULT_RULEBOOK)

# File processing
if uploaded_file and rulebook_text and instructions_text:
    with st.spinner("Extracting text from contract..."):
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == "pdf":
            contract_text = extract_text_from_pdf(uploaded_file)
        elif file_type == "docx":
            contract_text = extract_text_from_word(uploaded_file)
        else:
            st.error("Unsupported file type.")
            st.stop()

    text_chunks = split_text(contract_text, GEMINI_MAX_WORDS)

    if st.button("Analyze Risks"):
        with st.spinner("Analyzing contract risks..."):
            risk_analysis = ""
            for chunk in text_chunks:
                result = analyze_contract(chunk, rulebook_text, instructions_text)
                result = re.sub(r"~~(.*?)~~", r"<del>\1</del>", result)  # markdown strikethrough
                result = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", result)  # markdown bold
                risk_analysis += result + "\n\n"

            st.subheader("Risk Analysis Report")
            st.markdown(risk_analysis, unsafe_allow_html=True)

            # Convert Markdown to HTML
            risk_analysis_html = markdown.markdown(risk_analysis, extensions=["tables"])
            docx_bytes = report.report_downloader(risk_analysis_html, logo_path="fm-logo.png")

            st.download_button(
                label="Download Report as Word (.docx)",
                data=docx_bytes,
                file_name="Risk_Analysis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.info("Please upload a contract and ensure the rulebook + instructions are filled in.")
